# ============================================================================
# Trading Journal Web App v1.1
# Modern web interface with TradingView charts and Claude AI
# Cloud-ready deployment
# ============================================================================

from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from functools import wraps
from datetime import datetime, timezone, timedelta
import os
import json
import anthropic

# Optional imports
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    import yfinance as yf
except ImportError:
    yf = None

try:
    import cloudinary
    import cloudinary.uploader
except ImportError:
    cloudinary = None

import requests
import base64
import io

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# ============================================================================
# Configuration - Cloud Ready
# ============================================================================

# Detect if running in cloud (no local file system access)
IS_CLOUD = os.environ.get('IS_CLOUD', 'false').lower() == 'true'

# Password protection (set APP_PASSWORD env var to enable) - legacy single-user mode
APP_PASSWORD = os.environ.get('APP_PASSWORD', '')

# Multi-user mode: invite code for registration
INVITE_CODE = os.environ.get('INVITE_CODE', '')

# Users database path
USERS_DB_PATH = os.path.join(os.path.dirname(__file__), 'users.json')

# ============================================================================
# Authentication & User Management
# ============================================================================

import hashlib

def hash_password(password):
    """Simple password hashing"""
    return hashlib.sha256(password.encode()).hexdigest()

def load_users():
    """Load users database"""
    if os.path.exists(USERS_DB_PATH):
        try:
            with open(USERS_DB_PATH, 'r') as f:
                return json.load(f)
        except:
            pass
    return {}

def save_users(users):
    """Save users database"""
    with open(USERS_DB_PATH, 'w') as f:
        json.dump(users, f, indent=2)

def create_user(username, password):
    """Create a new user"""
    users = load_users()
    if username.lower() in users:
        return False, "Username already exists"
    users[username.lower()] = {
        'username': username,
        'password_hash': hash_password(password),
        'created': datetime.now().isoformat()
    }
    save_users(users)
    return True, "Account created"

def verify_user(username, password):
    """Verify user credentials"""
    users = load_users()
    user = users.get(username.lower())
    if user and user['password_hash'] == hash_password(password):
        return True, user['username']
    return False, None

def get_current_user():
    """Get current logged-in username"""
    return session.get('username', 'default')

def is_multi_user_mode():
    """Check if multi-user mode is enabled (invite code is set)"""
    return bool(INVITE_CODE)

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Multi-user mode (invite code set)
        if is_multi_user_mode():
            if not session.get('authenticated') or not session.get('username'):
                if request.path.startswith('/api/'):
                    return jsonify({'error': 'Authentication required'}), 401
                return redirect(url_for('login'))
            return f(*args, **kwargs)

        # Legacy single-user mode (APP_PASSWORD)
        if APP_PASSWORD:
            if not session.get('authenticated'):
                if request.path.startswith('/api/'):
                    return jsonify({'error': 'Authentication required'}), 401
                return redirect(url_for('login'))
            return f(*args, **kwargs)

        # No auth required
        return f(*args, **kwargs)
    return decorated_function

# Local paths (used when running locally)
JOURNAL_PATH = os.environ.get('JOURNAL_PATH', r"D:\OneDrive - Dick Koch LLC\Trading Journal.docx")
CONFIG_PATH = os.path.join(os.path.expanduser("~"), ".trading_journal_web_config")

# Cloud storage: use JSON file in app directory or environment variables
CLOUD_JOURNAL_PATH = os.path.join(os.path.dirname(__file__), 'journal_entries.json')

SYSTEM_PROMPT = """You are my Trading Journal. Each time I speak or ramble, convert my thoughts into a structured trading-journal entry. Extract and organize only what I actually say:

• My emotional state
• Market conditions I mentioned
• Trades I took or considered
• My reasoning and expectations
• Mistakes I noticed
• Lessons learned
• Follow-up items for the next session

Rules:
• Keep the tone neutral, factual, and concise.
• If I speak in fragments or ramble, infer structure but do not invent facts.
• Do not add commentary beyond what I explicitly said.
• Reorganize my thoughts into a clean, professional journal entry.

End each entry with:
• A short 2–3 sentence summary
• 1–3 reflection questions that would help me improve

Format the entry with clear section headers using bullet points. Start with a header line showing the date and time provided.

If market data is provided, incorporate relevant observations into the Market Conditions section."""


# ============================================================================
# Helper Functions
# ============================================================================

def get_est_timestamp():
    est = timezone(timedelta(hours=-5))
    return datetime.now(est).strftime('%A, %B %d, %Y %I:%M %p EST')

def load_config():
    """Load all config settings - environment variables take priority"""
    config = {
        "api_key": "",
        "cloudinary_cloud_name": "",
        "cloudinary_api_key": "",
        "cloudinary_api_secret": ""
    }

    # Environment variables (required for cloud deployment)
    env_mappings = {
        "ANTHROPIC_API_KEY": "api_key",
        "CLOUDINARY_CLOUD_NAME": "cloudinary_cloud_name",
        "CLOUDINARY_API_KEY": "cloudinary_api_key",
        "CLOUDINARY_API_SECRET": "cloudinary_api_secret"
    }

    for env_var, config_key in env_mappings.items():
        env_val = os.environ.get(env_var)
        if env_val:
            config[config_key] = env_val

    # Load from config file (local mode only, won't override env vars)
    if not IS_CLOUD and os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH) as f:
                data = json.load(f)
                # Only use file values if env vars are not set
                for key, val in data.items():
                    if not config.get(key):
                        config[key] = val
        except:
            pass

    return config

def save_config(config):
    """Save config settings (only works in local mode)"""
    if IS_CLOUD:
        # In cloud mode, settings come from environment variables
        # Return True to indicate "success" even though we don't save
        return True

    try:
        existing = load_config()
        existing.update(config)
        with open(CONFIG_PATH, "w") as f:
            json.dump(existing, f)
        return True
    except:
        return False

def get_api_key():
    return load_config().get("api_key", "")

def save_api_key(key):
    return save_config({"api_key": key})

def configure_cloudinary():
    """Configure Cloudinary with saved credentials"""
    if cloudinary is None:
        return False
    config = load_config()
    cloud_name = config.get("cloudinary_cloud_name", "")
    api_key = config.get("cloudinary_api_key", "")
    api_secret = config.get("cloudinary_api_secret", "")

    if cloud_name and api_key and api_secret:
        cloudinary.config(
            cloud_name=cloud_name,
            api_key=api_key,
            api_secret=api_secret
        )
        return True
    return False

# ============================================================================
# Cloud Journal Storage (JSON-based, user-specific)
# ============================================================================

def get_user_journal_path(username=None):
    """Get the journal path for a specific user"""
    if username is None:
        username = get_current_user()
    # Sanitize username for filename
    safe_username = "".join(c for c in username if c.isalnum() or c in '-_').lower()
    if not safe_username:
        safe_username = "default"
    return os.path.join(os.path.dirname(__file__), f'journal_{safe_username}.json')

def load_cloud_journal(username=None):
    """Load journal entries from JSON file (cloud mode, user-specific)"""
    journal_path = get_user_journal_path(username)
    if os.path.exists(journal_path):
        try:
            with open(journal_path, 'r') as f:
                return json.load(f)
        except:
            return {"entries": []}
    return {"entries": []}

def save_cloud_journal(data, username=None):
    """Save journal entries to JSON file (cloud mode, user-specific)"""
    journal_path = get_user_journal_path(username)
    try:
        with open(journal_path, 'w') as f:
            json.dump(data, f, indent=2)
        return True
    except Exception as e:
        print(f"Error saving cloud journal: {e}")
        return False

def add_cloud_entry(entry, username=None):
    """Add a new entry to the cloud journal"""
    journal = load_cloud_journal(username)
    journal["entries"].append(entry)
    return save_cloud_journal(journal, username)

def get_market_data():
    if yf is None:
        return None

    try:
        data = {}
        symbols = {
            '^GSPC': 'SPX',
            '^NDX': 'NDX',
            '^RUT': 'RUT',
            '^DJI': 'DJI',
            'GLD': 'Gold',
            'BTC-USD': 'Bitcoin',
            'TLT': 'TLT',
        }

        for sym, name in symbols.items():
            try:
                t = yf.Ticker(sym)
                # Use 5 day history with daily interval for reliable close-to-close data
                hist = t.history(period='5d', interval='1d')
                if len(hist) >= 2:
                    # Today's close (or most recent)
                    curr = float(hist['Close'].iloc[-1])
                    # Previous day's close
                    prev_close = float(hist['Close'].iloc[-2])
                    ch = ((curr - prev_close) / prev_close) * 100
                    data[sym] = {
                        'name': name,
                        'price': round(curr, 2),
                        'change': round(ch, 2),
                        'direction': 'up' if ch > 0 else 'down' if ch < 0 else 'flat'
                    }
            except:
                pass

        # VIX
        try:
            t = yf.Ticker('^VIX')
            hist = t.history(period='1d', interval='1m')
            if len(hist):
                curr = float(hist['Close'].iloc[-1])
                prev = t.info.get("previousClose", curr)
                ch = ((curr - prev) / prev) * 100 if prev else 0
                data['VIX'] = {
                    'name': 'VIX',
                    'price': round(curr, 2),
                    'change': round(ch, 2),
                    'direction': 'up' if ch > 0 else 'down' if ch < 0 else 'flat',
                    'status': 'LOW' if curr < 15 else 'ELEVATED' if curr > 25 else 'NORMAL'
                }
        except:
            pass

        # Calculate sentiment based on SPX
        spy_ch = data.get('^GSPC', {}).get('change', 0)
        vix = data.get('VIX', {}).get('price', 20)

        score = 0
        if spy_ch > 1: score += 2
        elif spy_ch > 0.3: score += 1
        elif spy_ch < -1: score -= 2
        elif spy_ch < -0.3: score -= 1

        if vix < 15: score += 1
        elif vix > 25: score -= 1
        if vix > 30: score -= 1

        if score >= 3:
            sentiment = {'label': 'STRONGLY BULLISH', 'color': 'green', 'icon': '🟢'}
        elif score >= 1:
            sentiment = {'label': 'BULLISH', 'color': 'lightgreen', 'icon': '🟢'}
        elif score <= -3:
            sentiment = {'label': 'STRONGLY BEARISH', 'color': 'red', 'icon': '🔴'}
        elif score <= -1:
            sentiment = {'label': 'BEARISH', 'color': 'salmon', 'icon': '🔴'}
        else:
            sentiment = {'label': 'NEUTRAL', 'color': 'yellow', 'icon': '🟡'}

        data['sentiment'] = sentiment
        data['timestamp'] = datetime.now(timezone(timedelta(hours=-5))).strftime('%I:%M:%S %p EST')

        return data
    except Exception as e:
        print(f"Market data error: {e}")
        return None

def format_market_for_prompt(data):
    if not data:
        return ""

    lines = ["Current Market Conditions:"]
    for sym in ['^GSPC', '^NDX', '^RUT', '^DJI', 'GLD', 'BTC-USD', 'TLT', 'VIX']:
        if sym in data:
            d = data[sym]
            sign = '+' if d['change'] > 0 else ''
            lines.append(f"  {d['name']} ({sym}): ${d['price']} ({sign}{d['change']}%)")

    if 'VIX' in data:
        v = data['VIX']
        sign = '+' if v['change'] > 0 else ''
        lines.append(f"  VIX: {v['price']} ({sign}{v['change']}%) - {v['status']}")

    if 'sentiment' in data:
        lines.append(f"  Overall Sentiment: {data['sentiment']['icon']} {data['sentiment']['label']}")

    return '\n'.join(lines)


# ============================================================================
# Routes
# ============================================================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # Multi-user mode
        if is_multi_user_mode():
            username = request.form.get('username', '').strip()
            password = request.form.get('password', '')
            success, display_name = verify_user(username, password)
            if success:
                session['authenticated'] = True
                session['username'] = display_name
                return redirect(url_for('index'))
            return render_template('login.html', error='Invalid username or password', multi_user=True)

        # Legacy single-user mode
        password = request.form.get('password', '')
        if password == APP_PASSWORD:
            session['authenticated'] = True
            session['username'] = 'default'
            return redirect(url_for('index'))
        return render_template('login.html', error='Invalid password', multi_user=False)

    return render_template('login.html', multi_user=is_multi_user_mode())

@app.route('/register', methods=['GET', 'POST'])
def register():
    # Only allow registration in multi-user mode
    if not is_multi_user_mode():
        return redirect(url_for('login'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        invite_code = request.form.get('invite_code', '')

        # Validate invite code
        if invite_code != INVITE_CODE:
            return render_template('register.html', error='Invalid invite code')

        # Validate username
        if not username or len(username) < 3:
            return render_template('register.html', error='Username must be at least 3 characters')

        if not username.replace('_', '').replace('-', '').isalnum():
            return render_template('register.html', error='Username can only contain letters, numbers, - and _')

        # Validate password
        if not password or len(password) < 4:
            return render_template('register.html', error='Password must be at least 4 characters')

        if password != confirm_password:
            return render_template('register.html', error='Passwords do not match')

        # Create user
        success, message = create_user(username, password)
        if success:
            # Auto-login after registration
            session['authenticated'] = True
            session['username'] = username
            return redirect(url_for('index'))
        return render_template('register.html', error=message)

    return render_template('register.html')

@app.route('/logout')
def logout():
    session.pop('authenticated', None)
    session.pop('username', None)
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    return render_template('index.html')

@app.route('/api/market-data')
@login_required
def api_market_data():
    data = get_market_data()
    if data:
        return jsonify(data)
    return jsonify({'error': 'Could not fetch market data'}), 500

@app.route('/api/process-entry', methods=['POST'])
@login_required
def api_process_entry():
    try:
        data = request.json
        raw_text = data.get('text', '')
        include_market = data.get('include_market', True)

        if not raw_text:
            return jsonify({'error': 'No text provided'}), 400

        api_key = get_api_key()
        if not api_key:
            return jsonify({'error': 'API key not configured'}), 400

        # Build message
        timestamp = get_est_timestamp()
        msg = f"Date/Time: {timestamp}\n\n"

        if include_market:
            market_data = get_market_data()
            if market_data:
                msg += format_market_for_prompt(market_data) + "\n\n"

        msg += f"My trading thoughts:\n{raw_text}"

        # Call Claude
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": msg}]
        )

        formatted_entry = response.content[0].text

        # Build full entry
        market_data = get_market_data()
        sentiment = market_data.get('sentiment', {}) if market_data else {}

        full_entry = {
            'timestamp': timestamp,
            'sentiment': sentiment,
            'content': formatted_entry,
            'raw_text': raw_text
        }

        return jsonify(full_entry)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-journal', methods=['POST'])
@login_required
def api_save_journal():
    try:
        data = request.json
        entries = data.get('entries', [])

        if not entries:
            return jsonify({'error': 'No entries to save'}), 400

        # Cloud mode: save to JSON
        if IS_CLOUD:
            for entry in entries:
                # Ensure entry has required fields
                cloud_entry = {
                    'timestamp': entry.get('timestamp', get_est_timestamp()),
                    'sentiment': entry.get('sentiment', {}),
                    'content': entry.get('content', ''),
                    'saved_at': datetime.now(timezone.utc).isoformat()
                }
                if not add_cloud_entry(cloud_entry):
                    return jsonify({'error': 'Failed to save entry'}), 500

            return jsonify({'success': True, 'count': len(entries), 'mode': 'cloud'})

        # Local mode: save to Word document
        if Document is None:
            return jsonify({'error': 'python-docx not installed'}), 500

        # Open or create document
        if os.path.exists(JOURNAL_PATH):
            doc = Document(JOURNAL_PATH)
        else:
            doc = Document()
            h = doc.add_heading("Trading Journal", 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add entries
        for entry in entries:
            if len(doc.paragraphs) > 1:
                doc.add_page_break()

            # Header
            p = doc.add_paragraph("═" * 60)
            p.style.font.name = 'Consolas'

            p = doc.add_paragraph("TRADING JOURNAL ENTRY")
            p.style.font.bold = True
            p.style.font.size = Pt(14)

            p = doc.add_paragraph(entry.get('timestamp', ''))

            sentiment = entry.get('sentiment', {})
            if sentiment:
                p = doc.add_paragraph(f"Market Sentiment: {sentiment.get('icon', '')} {sentiment.get('label', '')}")
                p.style.font.bold = True
                if 'BULLISH' in sentiment.get('label', ''):
                    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif 'BEARISH' in sentiment.get('label', ''):
                    p.runs[0].font.color.rgb = RGBColor(200, 0, 0)

            p = doc.add_paragraph("═" * 60)
            p.style.font.name = 'Consolas'

            # Content - process line by line, embedding images
            content = entry.get('content', '')
            for line in content.split('\n'):
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                # Check if this line contains an image URL from Cloudinary
                if line_stripped.startswith('[Image') and 'cloudinary.com' in line_stripped:
                    # Extract URL from format: [Image 1]: https://...
                    try:
                        url = line_stripped.split(']: ', 1)[1].strip()
                        # Download the image
                        img_response = requests.get(url, timeout=10)
                        if img_response.status_code == 200:
                            # Add image to document
                            img_stream = io.BytesIO(img_response.content)
                            doc.add_paragraph()  # Add spacing
                            p = doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(img_stream, width=Inches(5))
                            doc.add_paragraph()  # Add spacing after
                        else:
                            # If download fails, just add the URL as text
                            p = doc.add_paragraph(line)
                            p.style.font.name = 'Calibri'
                            p.style.font.size = Pt(11)
                    except Exception as img_error:
                        print(f"Failed to embed image: {img_error}")
                        p = doc.add_paragraph(line)
                        p.style.font.name = 'Calibri'
                        p.style.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(line)
                    p.style.font.name = 'Calibri'
                    p.style.font.size = Pt(11)

        doc.save(JOURNAL_PATH)

        return jsonify({'success': True, 'count': len(entries), 'path': JOURNAL_PATH})

    except PermissionError:
        return jsonify({'error': 'File is open in another program'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/settings', methods=['GET', 'POST'])
@login_required
def api_settings():
    if request.method == 'GET':
        config = load_config()

        # Check if API key came from environment variable
        api_key_from_env = bool(os.environ.get('ANTHROPIC_API_KEY'))

        return jsonify({
            'has_key': bool(config.get('api_key')),
            'key_preview': config.get('api_key', '')[:10] + '...' if config.get('api_key') else '',
            'key_from_env': api_key_from_env,
            'is_cloud': IS_CLOUD,
            'has_cloudinary': bool(config.get('cloudinary_cloud_name') and config.get('cloudinary_api_key') and config.get('cloudinary_api_secret')),
            'cloudinary_cloud_name': config.get('cloudinary_cloud_name', ''),
            'cloudinary_from_env': bool(os.environ.get('CLOUDINARY_CLOUD_NAME'))
        })
    else:
        data = request.json
        config_update = {}

        if 'api_key' in data:
            config_update['api_key'] = data['api_key']
        if 'cloudinary_cloud_name' in data:
            config_update['cloudinary_cloud_name'] = data['cloudinary_cloud_name']
        if 'cloudinary_api_key' in data:
            config_update['cloudinary_api_key'] = data['cloudinary_api_key']
        if 'cloudinary_api_secret' in data:
            config_update['cloudinary_api_secret'] = data['cloudinary_api_secret']

        if save_config(config_update):
            return jsonify({'success': True})
        return jsonify({'error': 'Could not save settings'}), 500

@app.route('/api/upload-image', methods=['POST'])
@login_required
def api_upload_image():
    """Upload an image to Cloudinary"""
    try:
        if cloudinary is None:
            return jsonify({'error': 'Cloudinary not installed'}), 500

        if not configure_cloudinary():
            return jsonify({'error': 'Cloudinary not configured. Please add credentials in Settings.'}), 400

        data = request.json
        image_data = data.get('image', '')

        if not image_data:
            return jsonify({'error': 'No image data provided'}), 400

        # Handle base64 data URL
        if image_data.startswith('data:'):
            # Extract base64 part
            image_data = image_data.split(',')[1]

        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        public_id = f"trading_journal/{timestamp}"

        # Upload to Cloudinary
        result = cloudinary.uploader.upload(
            f"data:image/png;base64,{image_data}",
            public_id=public_id,
            folder="trading_journal"
        )

        return jsonify({
            'success': True,
            'url': result['secure_url'],
            'public_id': result['public_id'],
            'width': result.get('width'),
            'height': result.get('height')
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/view-journal')
@login_required
def api_view_journal():
    try:
        # Cloud mode: read from JSON
        if IS_CLOUD:
            journal = load_cloud_journal()
            if not journal['entries']:
                return jsonify({'content': None, 'message': 'No journal entries found'})

            # Format entries for display
            content_parts = []
            for entry in journal['entries']:
                content_parts.append("═" * 60)
                content_parts.append("TRADING JOURNAL ENTRY")
                content_parts.append(entry.get('timestamp', ''))
                if entry.get('sentiment'):
                    s = entry['sentiment']
                    content_parts.append(f"Market Sentiment: {s.get('icon', '')} {s.get('label', '')}")
                content_parts.append("═" * 60)
                content_parts.append(entry.get('content', ''))
                content_parts.append("")

            return jsonify({'content': '\n'.join(content_parts)})

        # Local mode: read from Word document
        if not os.path.exists(JOURNAL_PATH):
            return jsonify({'content': None, 'message': 'No journal file found'})

        if Document is None:
            return jsonify({'error': 'python-docx not installed'}), 500

        doc = Document(JOURNAL_PATH)
        content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)

        return jsonify({'content': '\n'.join(content)})

    except PermissionError:
        return jsonify({'error': 'Journal file is open in another program'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/list-entries')
@login_required
def api_list_entries():
    """List the last 10 journal entries with previews"""
    try:
        # Cloud mode: read from JSON
        if IS_CLOUD:
            journal = load_cloud_journal()
            if not journal['entries']:
                return jsonify({'entries': [], 'message': 'No journal entries found'})

            entries = []
            for entry in journal['entries']:
                preview = entry.get('content', '')[:100] + '...' if entry.get('content') else ''
                entries.append({
                    'timestamp': entry.get('timestamp', ''),
                    'sentiment': entry.get('sentiment', {}).get('label', ''),
                    'preview': preview,
                    'content': entry.get('content', '')
                })

            # Return last 10 entries, most recent first
            entries.reverse()
            return jsonify({'entries': entries[:10]})

        # Local mode: read from Word document
        if not os.path.exists(JOURNAL_PATH):
            return jsonify({'entries': [], 'message': 'No journal file found'})

        if Document is None:
            return jsonify({'error': 'python-docx not installed'}), 500

        doc = Document(JOURNAL_PATH)

        # Parse entries - each entry starts with the separator line
        entries = []
        current_entry = []
        entry_timestamp = None
        entry_sentiment = None

        for para in doc.paragraphs:
            text = para.text.strip()

            # Detect entry separator
            if text.startswith('═' * 20) or text == 'TRADING JOURNAL ENTRY':
                if current_entry and entry_timestamp:
                    # Save previous entry
                    content = '\n'.join(current_entry)
                    # Get preview (first 100 chars of actual content)
                    preview_lines = [l for l in current_entry if l and not l.startswith('═') and l != 'TRADING JOURNAL ENTRY' and 'Market Sentiment' not in l]
                    preview = ' '.join(preview_lines)[:100] + '...' if preview_lines else ''

                    entries.append({
                        'timestamp': entry_timestamp,
                        'sentiment': entry_sentiment,
                        'preview': preview,
                        'content': content
                    })

                if text == 'TRADING JOURNAL ENTRY':
                    current_entry = [text]
                    entry_timestamp = None
                    entry_sentiment = None
                else:
                    current_entry = [text]
            elif current_entry is not None:
                current_entry.append(text)

                # Try to extract timestamp (usually line after TRADING JOURNAL ENTRY)
                if entry_timestamp is None and text and not text.startswith('═') and 'Market Sentiment' not in text:
                    # Check if it looks like a date
                    if any(day in text for day in ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']):
                        entry_timestamp = text

                # Extract sentiment
                if 'Market Sentiment:' in text:
                    entry_sentiment = text.replace('Market Sentiment:', '').strip()

        # Don't forget the last entry
        if current_entry and entry_timestamp:
            content = '\n'.join(current_entry)
            preview_lines = [l for l in current_entry if l and not l.startswith('═') and l != 'TRADING JOURNAL ENTRY' and 'Market Sentiment' not in l]
            preview = ' '.join(preview_lines)[:100] + '...' if preview_lines else ''

            entries.append({
                'timestamp': entry_timestamp,
                'sentiment': entry_sentiment,
                'preview': preview,
                'content': content
            })

        # Return last 10 entries, most recent first
        entries.reverse()
        return jsonify({'entries': entries[:10]})

    except PermissionError:
        return jsonify({'error': 'Journal file is open in another program'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update-entry', methods=['POST'])
@login_required
def api_update_entry():
    """Update an existing entry in the journal"""
    try:
        data = request.json
        original_timestamp = data.get('original_timestamp', '')
        new_content = data.get('content', '')

        if not original_timestamp or not new_content:
            return jsonify({'error': 'Missing timestamp or content'}), 400

        if Document is None:
            return jsonify({'error': 'python-docx not installed'}), 500

        if not os.path.exists(JOURNAL_PATH):
            return jsonify({'error': 'Journal file not found'}), 404

        doc = Document(JOURNAL_PATH)

        # Find and update the entry
        in_target_entry = False
        entry_start_idx = None
        entry_end_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if text == 'TRADING JOURNAL ENTRY':
                if in_target_entry:
                    entry_end_idx = i
                    break
                entry_start_idx = i

            if entry_start_idx is not None and original_timestamp in text:
                in_target_entry = True

        if not in_target_entry:
            return jsonify({'error': 'Entry not found'}), 404

        # For now, we'll append updates rather than modifying in place
        # This is safer and preserves history
        # The new content will be saved as a new entry when user clicks Save

        return jsonify({'success': True, 'message': 'Entry loaded for editing'})

    except PermissionError:
        return jsonify({'error': 'Journal file is open in another program'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-entry', methods=['POST'])
@login_required
def api_delete_entry():
    """Delete a specific entry by index"""
    try:
        data = request.json
        entry_index = data.get('index')

        if entry_index is None:
            return jsonify({'error': 'Missing entry index'}), 400

        if IS_CLOUD:
            journal = load_cloud_journal()
            entries = journal.get('entries', [])

            # Convert to positive index if negative
            if entry_index < 0:
                entry_index = len(entries) + entry_index

            if 0 <= entry_index < len(entries):
                deleted = entries.pop(entry_index)
                save_cloud_journal(journal)
                return jsonify({'success': True, 'deleted': deleted.get('timestamp', 'Unknown')})
            else:
                return jsonify({'error': 'Entry not found'}), 404
        else:
            return jsonify({'error': 'Delete not supported in local mode'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update-entry', methods=['POST'])
@login_required
def api_update_entry():
    """Update an existing entry by index"""
    try:
        data = request.json
        entry_index = data.get('index')
        new_content = data.get('content')

        if entry_index is None:
            return jsonify({'error': 'Missing entry index'}), 400
        if not new_content:
            return jsonify({'error': 'Missing content'}), 400

        if IS_CLOUD:
            journal = load_cloud_journal()
            entries = journal.get('entries', [])

            if 0 <= entry_index < len(entries):
                # Update the entry content
                entries[entry_index]['content'] = new_content
                # Update timestamp to show it was edited
                entries[entry_index]['edited'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                # Update preview
                preview_text = new_content[:150].replace('\n', ' ')
                if len(new_content) > 150:
                    preview_text += '...'
                entries[entry_index]['preview'] = preview_text

                save_cloud_journal(journal)
                return jsonify({'success': True, 'message': 'Entry updated'})
            else:
                return jsonify({'error': 'Entry not found'}), 404
        else:
            return jsonify({'error': 'Update not supported in local mode'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear-journal', methods=['POST'])
@login_required
def api_clear_journal():
    """Clear all journal entries (cloud mode only)"""
    try:
        if IS_CLOUD:
            save_cloud_journal({'entries': []})
            return jsonify({'success': True, 'message': 'Journal cleared'})
        else:
            return jsonify({'error': 'Clear not supported in local mode - delete the Word file manually'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# Main
# ============================================================================

if __name__ == '__main__':
    # Get port from environment (for cloud platforms) or default to 5000
    port = int(os.environ.get('PORT', 5000))
    debug = not IS_CLOUD  # Disable debug mode in production

    print("Starting Trading Journal Web App...")
    if IS_CLOUD:
        print(f"Running in CLOUD mode on port {port}")
    else:
        print(f"Open http://localhost:{port} in your browser")

    app.run(host='0.0.0.0', port=port, debug=debug)
